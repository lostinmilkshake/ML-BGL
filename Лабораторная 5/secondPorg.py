from docx import Document
wordDoc = Document('БурковМП ГарцевЕА ЛитвиновКЛ 7301 лаб 5.docx')
u = []
for i in range(1, len(wordDoc.tables[1].rows)):
    u.append(float(wordDoc.tables[1].rows[i].cells[1].text))
ucp = sum(u) / len(u)
s2u = 0
for uElement in u:
    s2u += (uElement - ucp)**2
s2u = s2u * 1/(len(u) - 1)
s2ucp = s2u / len(u)
wordDoc.tables[1].rows[1].cells[2].text = str(s2u) + '\n' + str(s2u**0.5)
wordDoc.tables[1].rows[1].cells[3].text = str(s2ucp) + '\n' + str(s2ucp**0.5)
# TODO почситать доверительный интервал
wordDoc.save('БурковМП ГарцевЕА ЛитвиновКЛ 7301 лаб 5.docx')
