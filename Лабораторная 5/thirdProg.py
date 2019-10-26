from docx import Document
wordDoc = Document('БурковМП ГарцевЕА ЛитвиновКЛ 7301 лаб 5.docx')
u1 = []
u2 = []
for i in range(1, len(wordDoc.tables[2].rows)):
    u1.append(float(wordDoc.tables[2].rows[i].cells[1].text))
    u2.append(float(wordDoc.tables[3].rows[i].cells[1].text))
u1cp = sum(u1) / len(u1)
u2cp = sum(u2) / len(u2)
s2u1 = 0
s2u2 = 0
for uElement in u1:
    s2u1 += (uElement - u1cp)**2
for uElement in u2:
    s2u1 += (uElement - u1cp)**2
s2u1 = s2u1 * 1/(len(u1) - 1)
s2u2 = s2u2 * 1/(len(u2) - 1)
s2u1cp = s2u1 / len(u1)
s2u2cp = s2u2 / len(u2)
wordDoc.tables[2].rows[1].cells[2].text = str(s2u1) + '\n' + str(s2u1**0.5)
wordDoc.tables[2].rows[1].cells[3].text = str(s2u1cp) + '\n' + str(s2u1cp**0.5)
wordDoc.tables[3].rows[1].cells[2].text = str(s2u2) + '\n' + str(s2u2**0.5)
wordDoc.tables[3].rows[1].cells[3].text = str(s2u2cp) + '\n' + str(s2u2cp**0.5)
# TODO почситать доверительный интервал
r0 = 1000  # TODO проверить R0
p = u1cp*u2cp / r0
print("Мощность: ", p)
# Доверительный интервал
s2p = 1 / (r0 ** 2) * (((u2cp / r0)**2) * s2u1cp + ((u1cp / r0)**2) * s2u2cp)
print("Доверительный интервал мощности: ", s2p)
# TODO записать итоговую погрешность
wordDoc.save('БурковМП ГарцевЕА ЛитвиновКЛ 7301 лаб 5.docx')
