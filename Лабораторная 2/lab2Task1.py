from matplotlib import *
from math import *
import matplotlib.pyplot as plt
from docx import Document
wordDoc = Document('БурковМП ГарцевЕА ЛитвиновКЛ 7301 лаб 2.docx')
#n = int(input("Ammount of measurements:"))
n = len(wordDoc.tables[1].rows)
U = []
Uyv = []
dU = [] #Should be something, idk
Uym = []
dUyv = []
dUym = []
Un = float(1) # or 3, depending on what will be the given number
delta = []
gamma = []
H = []
for i in range(n - 4):
    '''U.append(float(input()));
    Uyv.append(float(input()))
    Uym.append(float(input()))'''
    U.append(float(wordDoc.tables[1].rows[i + 3].cells[0].text))
    Uyv.append(float(wordDoc.tables[1].rows[i + 3].cells[1].text))
    Uym.append(float(wordDoc.tables[1].rows[i + 3].cells[2].text))
    #print("U = ", U, "Uyv = ", Uyv, "Uym", Uym)
    dUyv.append(U[i] - Uyv[i])
    wordDoc.tables[1].rows[i + 3].cells[3].text = str(dUyv[i])
    dUym.append(U[i] - Uym[i])
    wordDoc.tables[1].rows[i + 3].cells[4].text = str(dUym[i])
    dU.append(dUyv[i] + dUym[i]) #maybemaybe
    delta.append(100 * (dU[i] / U[i]))
    wordDoc.tables[1].rows[i + 3].cells[5].text = str(delta[i])
    gamma.append(100 * (dU[i] / Un))
    wordDoc.tables[1].rows[i + 3].cells[6].text = str(gamma[i])
    H.append(100 * (abs(Uyv[i] - Uym[i]) / Un))
    wordDoc.tables[1].rows[i + 3].cells[7].text = str(H[i])
print("Gamma max = ", max(gamma), "Hmax = ", max(H))
plt.plot(U, delta, U, H)
plt.xlabel('U, В')
plt.ylabel(r'$\delta, H$')
plt.grid(True)
plt.show()
print("dUyv = ", dUyv, "dUym = ", dUym)
print("delta = ", delta, "gamma = ", gamma, "H = ", H)
wordDoc.save('БурковМП ГарцевЕА ЛитвиновКЛ 7301 лаб 2.docx')
