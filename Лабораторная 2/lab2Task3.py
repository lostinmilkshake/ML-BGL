from docx import Document
wordDoc = Document('БурковМП ГарцевЕА ЛитвиновКЛ 7301 лаб 2.docx')
Up = []
Ucp = []
U = []
delta = []
kf = [1.11, 1, 1.15]
for i in range(3):
    print(wordDoc.tables[4].columns[i + 1].cells[2].text)
    Up.append(float(wordDoc.tables[4].columns[i + 1].cells[2].text))
    Ucp.append(Up[i] / 1.11)
    wordDoc.tables[4].columns[i + 1].cells[3].text = str(Ucp[i])
    U.append(kf[i] * Ucp[i])
    wordDoc.tables[4].columns[i + 1].cells[4].text = str(U[i])
    delta.append(100 * (Up[i] - U[i]) / U[i])
    wordDoc.tables[4].columns[i + 1].cells[5].text = str(delta[i])
print(Ucp)
print(U)
print(delta)
wordDoc.save('БурковМП ГарцевЕА ЛитвиновКЛ 7301 лаб 2.docx')
