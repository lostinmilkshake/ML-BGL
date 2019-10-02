from matplotlib import *
from math import *
import matplotlib.pyplot as plt
from docx import Document
wordDoc = Document('БурковМП ГарцевЕА ЛитвиновКЛ 7301 лаб 2.docx')
n = len(wordDoc.tables[2].columns)
f = []
U = []
K = []
for i in range(n - 2):
    #For electric
    print(wordDoc.tables[2].columns[i + 2].cells[1].text)
    #f.append(float(wordDoc.tables[2].columns[i + 2].cells[1].text))
    #U.append(float(wordDoc.tables[2].columns[i + 2].cells[2].text))
    #K.append(U[i] / U[0])
    #wordDoc.tables[2].columns[i + 2].cells[3].text = str(K[i])
    #For digital
    #f.append(float(wordDoc.tables[2].columns[i + 2].cells[1].text))
    #U.append(float(wordDoc.tables[2].columns[i + 2].cells[4].text))
    #K.append(U[i] / U[0])
    #wordDoc.tables[2].columns[i + 2].cells[5].text = str(K[i])
    #For electric low
    #f.append(float(wordDoc.tables[3].columns[i + 2].cells[1].text))
    #U.append(float(wordDoc.tables[3].columns[i + 2].cells[2].text))
    #K.append(U[i] / U[0])
    #wordDoc.tables[3].columns[i + 2].cells[3].text = str(K[i])
    #For digital low
    #f.append(float(wordDoc.tables[3].columns[i + 2].cells[1].text))
    #U.append(float(wordDoc.tables[3].columns[i + 2].cells[4].text))
    #K.append(U[i] / U[0])
    #wordDoc.tables[3].columns[i + 2].cells[5].text = str(K[i])
plt.plot(f, K)
plt.grid(True)
plt.xlabel('f, кГц') #For low just гц
plt.ylabel('K(f)')
plt.show()
wordDoc.save('БурковМП ГарцевЕА ЛитвиновКЛ 7301 лаб 2.docx')
#find out what Fв/Fн should really be
