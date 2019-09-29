from docx import Document
wordDoc = Document('БурковМП ГарцевЕА ЛитвиновКЛ 7301 лаб 2.docx')
table_size = len(wordDoc.tables[2].columns)
cell = wordDoc.tables[4].columns[3].cells[2].text
print(cell)
