from docx import Document
from matplotlib import *
import matplotlib.pyplot as plt

wordDoc = Document('БурковМП_ГарцевЕА_ЛитвиновКЛ_7301_лаб_3.docx')
Rp = []
R = []
dR = []
for i in range(1, len(wordDoc.tables[1].rows)):
    Rp.append(float(wordDoc.tables[1].rows[i].cells[1].text))
    R.append(float(wordDoc.tables[1].rows[i].cells[2].text))
    dR.append(Rp[i-1] - R[i-1])
#График зависимости Rп от R
fig = plt.figure(1)
plt.plot(R, Rp)
plt.xlabel('R')
plt.ylabel(r'$R_п$')
plt.grid(True)
#График зависимости погрешности

secondPlt = plt.figure(2)
plt.plot(R, dR)
#plt.stairs(R, dR)
plt.grid(True)

plt.xlabel('R')
plt.ylabel(r'$\Delta R$')
plt.show()
fig.savefig('firstPlot.png')
secondPlt.savefig('secondPlot.png')
#wordDoc.save('БурковМП_ГарцевЕА_ЛитвиновКЛ_7301_лаб_3.docx')
