from docx import Document
from matplotlib import *
import math
import matplotlib.pyplot as plt

wordDoc = Document('БурковМП_ГарцевЕА_ЛитвиновКЛ_7301_лаб_3.docx')
RpN = []
RN = []
dRiN = []
q = 0.0001 #Высчетанное значение
for i in range(1, len(wordDoc.tables[2].rows)):
    RpN.append(float(wordDoc.tables[2].rows[i].cells[1].text))
    RN.append(float(wordDoc.tables[2].rows[i].cells[2].text))
    dRiN.append(RpN[i-1] - 0.5*q - RN[i - 1])
    wordDoc.tables[2].rows[i].cells[3].text = str(round(dRiN[i - 1], 3))
newFig = plt.figure()
plt.plot(RN, dRiN)
plt.grid(True)
#plt.xlabel(r'$\Delta R_{иN}')

plt.xlabel(r'$ R_N $') 
plt.ylabel(r'$\Delta R_{иN}$')
plt.show()
newFig.savefig('thirdPlot.png')

wordDoc.save('БурковМП_ГарцевЕА_ЛитвиновКЛ_7301_лаб_3.docx')
